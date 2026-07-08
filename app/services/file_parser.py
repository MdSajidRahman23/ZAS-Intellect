from pathlib import Path
import re
from typing import Tuple

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".py", ".js", ".java", ".cpp", ".c", ".html", ".css", ".md"}
WORD_PATTERN = re.compile(r"[A-Za-z0-9_]+|[\u0980-\u09FF]+")


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return clean_text("\n".join(chunks))


def parse_docx(path: Path) -> str:
    doc = Document(str(path))
    chunks = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    return clean_text("\n".join(chunks))


def parse_plain(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return clean_text(path.read_text(encoding=encoding, errors="ignore"))
        except Exception:
            continue
    return ""


def parse_submission(path: Path, max_chars: int = 50000) -> Tuple[str, str, int]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    if suffix == ".pdf":
        text = parse_pdf(path)
    elif suffix == ".docx":
        text = parse_docx(path)
    else:
        text = parse_plain(path)

    text = text[:max_chars]
    word_count = len(WORD_PATTERN.findall(text))
    return text, suffix.lstrip("."), word_count
